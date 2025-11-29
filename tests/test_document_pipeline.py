from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as WordDocument
from fpdf import FPDF

from mock_project.config import Settings
from mock_project.document_loader import load_documents, split_documents


def _create_docx(path: Path, text: str) -> None:
    doc = WordDocument()
    doc.add_heading("Mock Product Guide", level=1)
    doc.add_paragraph(text)
    doc.save(path)


def _create_pdf(path: Path, text: str) -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=text)
    pdf.output(path)


def _settings_for(tmp_path: Path) -> Settings:
    return Settings(
        openai_api_key="sk-test",
        chat_model="gpt-test",
        embedding_model="text-embedding-test",
        docs_path=tmp_path,
        chunk_size=200,
        chunk_overlap=50,
    )


@pytest.fixture()
def populated_docs(tmp_path: Path) -> Settings:
    _create_docx(tmp_path / "product_overview.docx", "Sản phẩm A giải quyết nhu cầu khách hàng doanh nghiệp.")
    _create_pdf(tmp_path / "service_catalog.pdf", "Gói dịch vụ Premium bao gồm hỗ trợ 24/7.")
    return _settings_for(tmp_path)


def test_load_documents_reads_pdf_and_docx(populated_docs: Settings) -> None:
    documents = load_documents(populated_docs)
    titles = {doc.metadata.get("source") for doc in documents}

    assert any("product_overview.docx" in title for title in titles)
    assert any("service_catalog.pdf" in title for title in titles)


def test_split_documents_creates_chunks(populated_docs: Settings) -> None:
    documents = load_documents(populated_docs)
    chunks = split_documents(populated_docs, documents)

    assert len(chunks) >= len(documents)
    assert all("Sản phẩm" in chunk.page_content or "Gói dịch vụ" in chunk.page_content for chunk in chunks)


