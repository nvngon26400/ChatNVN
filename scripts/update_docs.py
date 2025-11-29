from __future__ import annotations

from pathlib import Path
from unicodedata import normalize

from docx import Document
from fpdf import FPDF

DOCX_PATH = Path("data/docs/product_faq.docx")
PDF_PATH = Path("data/docs/product_faq.pdf")

CONTENT = {
    "title": "Product FAQ & Service Catalog",
    "intro": "Tài liệu tóm tắt danh mục sản phẩm và các dịch vụ hỗ trợ thông minh dành cho khách hàng doanh nghiệp (cập nhật 11/2025).",
    "sections": [
        (
            "Gói sản phẩm",
            [
                "1. Premium Suite: Phần mềm chăm sóc khách hàng đa kênh, chatbot AI và dashboard dữ liệu thời gian thực cho đội vận hành lớn.",
                "2. Growth Suite: Bộ automation marketing, chiến dịch nuôi dưỡng khách hàng và báo cáo đa kênh cho đội tăng trưởng.",
                "3. AI Copilot Add-on: Tự động gợi ý câu trả lời, phân loại ticket, phân tích cảm xúc và đề xuất knowledge base cần cập nhật.",
            ],
        ),
        (
            "Dịch vụ hỗ trợ",
            [
                "- Hỗ trợ 24/7 qua chat, hotline, email cho khách hàng Premium.",
                "- Tư vấn kỹ thuật chuyên sâu: solution architect đồng hành khi tích hợp API, tối ưu bảo mật và hiệu năng.",
                "- Success Manager cá nhân: thiết lập SLA, báo cáo sức khỏe và tổ chức workshop đào tạo nội bộ.",
                "- Health Check hàng quý: rà soát quy trình, đề xuất tối ưu chi phí và kế hoạch mở rộng.",
            ],
        ),
        (
            "Chính sách & quyền lợi",
            [
                "• Đổi trả 30 ngày cho mọi sản phẩm (điều kiện: chưa kích hoạt vĩnh viễn hoặc phần cứng không hư hại).",
                "• Sandbox miễn phí 60 ngày để thử tính năng mới.",
                "• Thông báo nâng cấp phần mềm trước 14 ngày, hỗ trợ rollback khi cần.",
            ],
        ),
        (
            "Dịch vụ thông minh",
            [
                "1. Smart Routing: định tuyến ticket theo ưu tiên và kỹ năng agent, giảm 35% thời gian phản hồi.",
                "2. Predictive Churn Alert: cảnh báo khách hàng có nguy cơ rời bỏ dựa trên hành vi sử dụng và điểm NPS.",
                "3. Knowledge Booster: tự động tóm tắt ca hỗ trợ thành bài viết và gợi ý nội dung cần cập nhật.",
                "4. Automation Builder: kéo-thả workflow để gửi survey, kích hoạt cảnh báo hoặc đồng bộ CRM khi ticket đóng.",
            ],
        ),
        (
            "Liên hệ nhanh",
            [
                "• Email: support@example.com",
                "• Hotline: 1900-123-456 (nhánh 2 cho Premium)",
                "• Portal: https://customer.example.com",
                "• SLA: phản hồi <30 phút (Premium) / <4 giờ (Growth).",
            ],
        ),
    ],
}


def main() -> None:
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    _build_docx(DOCX_PATH)
    _build_pdf(PDF_PATH)
    print(f"Updated knowledge base files: {DOCX_PATH.name}, {PDF_PATH.name}")


def _build_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading(CONTENT["title"], level=1)
    doc.add_paragraph(CONTENT["intro"])

    for heading, paragraphs in CONTENT["sections"]:
        doc.add_heading(heading, level=2)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)

    doc.save(path)


def _build_pdf(path: Path) -> None:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", "B", 16)
    pdf.multi_cell(0, 10, _ascii(CONTENT["title"]))
    pdf.ln(2)

    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 8, _ascii(CONTENT["intro"]))
    pdf.ln(4)

    for heading, paragraphs in CONTENT["sections"]:
        pdf.set_font("Arial", "B", 13)
        pdf.multi_cell(0, 8, _ascii(heading))
        pdf.ln(1)
        pdf.set_font("Arial", "", 12)
        for paragraph in paragraphs:
            pdf.multi_cell(0, 7, _ascii(paragraph))
            pdf.ln(0.5)
        pdf.ln(2)

    pdf.output(str(path))


def _ascii(text: str) -> str:
    """Ensure text is safe for standard PDF fonts."""

    return normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")


if __name__ == "__main__":
    main()
